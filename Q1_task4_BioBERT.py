import spacy  
from transformers import pipeline, AutoModelForTokenClassification, AutoTokenizer  
from collections import Counter  
from docx import Document  
import nltk  
from nltk.corpus import stopwords 

# Load SpaCy models  
#nlp_sci = spacy.load('en_core_sci_sm')  
#nlp_bc5cdr = spacy.load('en_ner_bc5cdr_md')  

# Load BioBERT model  
model_path = 'D:/BioBERT'  
model = AutoModelForTokenClassification.from_pretrained(model_path)  
tokenizer = AutoTokenizer.from_pretrained(model_path)  
ner_biobert = pipeline('ner', model=model, tokenizer=tokenizer)  

def extract_entities(text_file):  
    # Read text from file  
    with open(text_file, 'r', encoding='utf-8') as file:  
        text = file.read()  

    entities_biobert = []  

    # Use tokenizer to split text into valid chunks  
    # Tokenizer's max_length is generally 512, so we adjust for that  
    inputs = tokenizer(text, return_tensors='pt', truncation=True, padding=True, max_length=512, add_special_tokens=False)  
    tokens = inputs.input_ids.squeeze(0)  
    chunk_size = 512  # account for special tokens  

    for start in range(0, len(tokens), chunk_size):  
        end = min(start + chunk_size, len(tokens))  
        chunk_text = tokenizer.decode(tokens[start:end], skip_special_tokens=True)  
        
        words = nltk.word_tokenize(chunk_text.lower())
        stop_words = set(stopwords.words('english'))  

        # 去除停用词和非字母字符  
        filtered_words = [word for word in words if word.isalpha() and word not in stop_words]

        # Process with BioBERT  
        entities_biobert.extend([e for e in ner_biobert(filtered_words)])  

    return entities_biobert    

def analyze_and_output_results(text_file, output_word_file):  
    entities_biobert  = extract_entities(text_file)  
    
    # Count total entities  
    total_biobert = len(entities_biobert)  

    # Find most common entities  
    words = [ent['word'] for sublist in entities_biobert for ent in sublist]   
    common_biobert = Counter(words).most_common(5)  
    
    # Find differences  
    set_biobert = set(words)  


    # Create a Word document  
    doc = Document()  
    doc.add_heading('Entity Extraction and Comparison Results', level=1)  

    doc.add_heading('Total Entities Detected', level=2)  
    #doc.add_paragraph(f'SciSpacy: {total_sci}')  
    #doc.add_paragraph(f'BC5CDR: {total_cdr}')  
    doc.add_paragraph(f'BioBERT: {total_biobert}')  

    doc.add_heading('Most Common Entities', level=2)  
    #doc.add_paragraph(f'SciSpacy: {common_sci}')  
    #doc.add_paragraph(f'BC5CDR: {common_cdr}')  
    doc.add_paragraph(f'BioBERT: {common_biobert}')  

    doc.add_heading('Unique Entities Detected', level=2)  
    #doc.add_paragraph(f'SciSpacy: {diff_sci}')  
    #doc.add_paragraph(f'BC5CDR: {diff_cdr}')  
    doc.add_paragraph(f'BioBERT: {set_biobert}')  

    # Save the document to a file  
    doc.save(output_word_file)  
    print(f"Results have been saved to {output_word_file}")  

# Execute the improved function with the desired text file and output file path  
analyze_and_output_results('D:/project1/continue/test1/output/output1.txt', 'D:/project1/continue/test1/output/Entity_Comparison_Results_BIO.docx')
