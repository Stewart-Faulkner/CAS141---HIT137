import spacy  
from transformers import pipeline, AutoModelForTokenClassification, AutoTokenizer  
from collections import Counter  
from docx import Document  

# Load SpaCy models  
nlp_sci = spacy.load('en_core_sci_sm')  
nlp_bc5cdr = spacy.load('en_ner_bc5cdr_md')  


def extract_entities(text_file):  
    # Read text from file  
    with open(text_file, 'r', encoding='utf-8') as file:  
        text = file.read()  

    chunk_size = 1000000  
    chunks = [text[i:i + chunk_size] for i in range(0, len(text), chunk_size)]  

    entities_sci = []  
    entities_cdr = []  
   

    for chunk in chunks:  
        doc_sci = nlp_sci(chunk)  
        doc_cdr = nlp_bc5cdr(chunk)  

        entities_sci.extend([(ent.text, ent.label_) for ent in doc_sci.ents ])  
        entities_cdr.extend([(ent.text, ent.label_) for ent in doc_cdr.ents if ent.label_ in ['DISEASE', 'DRUG']])  
       
    return entities_sci, entities_cdr

def analyze_and_output_results(text_file, output_word_file):  
    entities_sci, entities_cdr = extract_entities(text_file)  
    
    # Count total entities  
    total_sci = len(entities_sci)  
    total_cdr = len(entities_cdr)  
  
    
    # Find most common entities  
    common_sci = Counter([ent[0] for ent in entities_sci]).most_common(5)  
    common_cdr = Counter([ent[0] for ent in entities_cdr]).most_common(5)  
  
    # Find differences  
    set_sci = set(ent[0] for ent in entities_sci)  
    set_cdr = set(ent[0] for ent in entities_cdr)  
  
    #diff_sci = set_sci - set_cdr - set_biobert  
    #diff_cdr = set_cdr - set_sci - set_biobert  
  
    # Create a Word document  
    doc = Document()  
    doc.add_heading('Entity Extraction and Comparison Results', level=1)  

    doc.add_heading('Total Entities Detected', level=2)  
    doc.add_paragraph(f'SciSpacy: {total_sci}')  
    doc.add_paragraph(f'BC5CDR: {total_cdr}')  


    doc.add_heading('Most Common Entities', level=2)  
    doc.add_paragraph(f'SciSpacy: {common_sci}')  
    doc.add_paragraph(f'BC5CDR: {common_cdr}')  

    doc.add_heading('Unique Entities Detected', level=2)  
    doc.add_paragraph(f'SciSpacy: {set_sci}')  
    doc.add_paragraph(f'BC5CDR: {set_cdr}')  

    # Save the document to a file  
    doc.save(output_word_file)  
    print(f"Results have been saved to {output_word_file}")  

# Execute the improved function with the desired text file and output file path  
analyze_and_output_results('D:/project1/continue/test1/output/output1.txt', 'D:/project1/continue/test1/output/Entity_Comparison_Results.docx')
